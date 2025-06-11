import os
from pathlib import Path
from typing import Dict, List, Any
import docx
from docx.shared import Pt, RGBColor
import json
from dotenv import load_dotenv
import google.generativeai as genai

class DocumentParser:
    def __init__(self):
        self.gemini_client = None
        self._init_gemini()
    
    def _init_gemini(self):
        load_dotenv()
        api_key = os.getenv('GEMINI_API_KEY')
        if api_key:
            genai.configure(api_key=api_key)
            self.gemini_client = genai.GenerativeModel('gemini-2.0-flash')
    
    def _init_document_styles(self, doc: docx.Document):
        styles = doc.styles
        
        # Title style
        title_style = styles['Title']
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 0)
        
        # Heading 1 style
        h1_style = styles['Heading 1']
        h1_font = h1_style.font
        h1_font.name = 'Arial'
        h1_font.size = Pt(18)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 0, 0)
        
        # Heading 2 style
        h2_style = styles['Heading 2']
        h2_font = h2_style.font
        h2_font.name = 'Arial'
        h2_font.size = Pt(16)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0, 0, 0)
        
        # Heading 3 style
        h3_style = styles['Heading 3']
        h3_font = h3_style.font
        h3_font.name = 'Arial'
        h3_font.size = Pt(14)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(0, 0, 0)
        
        # Normal text style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)
        normal_font.color.rgb = RGBColor(0, 0, 0)
    
    def _get_safe_style(self, doc: docx.Document, style_name: str) -> docx.styles.style._ParagraphStyle:
        try:
            return doc.styles[style_name]
        except KeyError:
            return doc.styles['Normal']
    
    def process_document_set(self, doc_set: Dict[str, List[str]], input_dir: Path) -> Dict[str, Any]:
        set_name = doc_set['name']
        documents = doc_set['documents']
        document_analyses = []
        
        for doc_name in documents:
            doc_path = input_dir / doc_name
            try:
                print(f"\nProcessing document: {doc_name}")
                doc = docx.Document(doc_path)
                text_content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                doc_analysis = self.analyze_document(text_content, doc_name)
                document_analyses.append(doc_analysis)
            except Exception as e:
                print(f"Error processing {doc_name}: {e}")
        
        set_summary = self.generate_comprehensive_summary(document_analyses)
        return {
            "set_name": set_name,
            "summary": set_summary,
            "documents": documents
        }
    
    def analyze_document(self, text: str, doc_name: str) -> Dict[str, Any]:
        try:
            prompt = f"""
            Analyze the following document and provide:
            1. Main topic and purpose
            2. Key points and findings
            3. Important context and background
            4. Critical information
            5. Recommendations or action items
            
            Document: {text}
            
            Format the response with clear sections and use **bold** for important terms.
            """
            
            if self.gemini_client:
                response = self.gemini_client.generate_content(prompt)
                analysis = response.text
            else:
                analysis = "LLM not configured"
            
            return {
                "document_name": doc_name,
                "analysis": analysis
            }
        except Exception as e:
            print(f"Error analyzing document {doc_name}: {e}")
            return {
                "document_name": doc_name,
                "analysis": "Error analyzing document"
            }
    
    def generate_comprehensive_summary(self, document_analyses: List[Dict[str, Any]]) -> str:
        try:
            prompt = f"""
            Create a comprehensive summary of the following documents, incorporating their individual analyses.
            
            Document Analyses:
            {json.dumps(document_analyses, indent=2)}
            
            Please provide a summary that includes:
            1. Executive Summary
            2. Key Findings from Each Document
            3. Important Context and Background
            4. Critical Information
            5. Cross-Document Insights
            6. Recommendations and Action Items
            
            Format the response with clear sections and use **bold** for important terms and headings.
            """
            
            if self.gemini_client:
                response = self.gemini_client.generate_content(prompt)
                summary = response.text
            else:
                summary = "LLM not configured"
            
            return summary
        except Exception as e:
            print(f"Error generating comprehensive summary: {e}")
            return "Error generating summary"
    
    def create_context_document(self, document_sets: List[Dict[str, List[str]]], input_dir: Path, output_file: str) -> docx.Document:
        doc = docx.Document()
        self._init_document_styles(doc)
        
        title = doc.add_paragraph()
        title.style = self._get_safe_style(doc, 'Heading 1')
        title.add_run(f"Document Set Analysis: {output_file}")
        
        toc = doc.add_paragraph()
        toc.style = self._get_safe_style(doc, 'Heading 2')
        toc.add_run("Table of Contents")
        
        for i, doc_set in enumerate(document_sets, 1):
            toc_item = doc.add_paragraph()
            toc_item.style = self._get_safe_style(doc, 'Normal')
            toc_item.add_run(f"{i}. {doc_set['name']}")
            
            set_heading = doc.add_paragraph()
            set_heading.style = self._get_safe_style(doc, 'Heading 2')
            set_heading.add_run(f"{i}. {doc_set['name']}")
            
            set_info = self.process_document_set(doc_set, input_dir)
            
            summary_heading = doc.add_paragraph()
            summary_heading.style = self._get_safe_style(doc, 'Heading 3')
            summary_heading.add_run("Summary")
            
            summary_text = set_info['summary']
            current_para = doc.add_paragraph()
            current_para.style = self._get_safe_style(doc, 'Normal')
            
            paragraphs = summary_text.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    parts = para_text.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 1:
                            current_para.add_run(part).bold = True
                        else:
                            current_para.add_run(part)
                    
                    current_para = doc.add_paragraph()
                    current_para.style = self._get_safe_style(doc, 'Normal')
            
            doc.add_paragraph("---" * 20, style=self._get_safe_style(doc, 'Normal'))
        
        return doc