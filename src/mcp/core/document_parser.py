from pathlib import Path
from typing import Dict, List, Any
from docx import Document
import logging
import markdown2
from html.parser import HTMLParser
from .llm_client import LLMClient
from .config import Settings
from .style_manager import StyleManager

logger = logging.getLogger(__name__)

class DocumentParser:
    def __init__(self):
        self.settings = Settings()
        self.llm_client = LLMClient(settings=self.settings)
        self.chunk_size = 4000
        self.chunk_overlap = 200
        self.style_manager = StyleManager()

    # Error handling
    def _handle_error(
        self, operation: str, doc_name: str, error: Exception
    ) -> Dict[str, Any]:
        """Handle errors during document processing."""
        return {
            "document_name": doc_name,
            "error": f"Error {operation} {doc_name}: {str(error)}",
        }

    # Class to parse HTML into a docx document
    class DocxHTMLParser(HTMLParser):
        def __init__(self, doc, style_map):
            super().__init__()
            self.doc = doc
            self.style_map = style_map
            self.current_para = None
            self.current_style = "Normal"
            self.bold = False
            self.italic = False
            self.list_type = None
            self.list_level = 0

        def handle_starttag(self, tag, attrs):
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                self.current_style = self.style_map.get(tag, "Normal")
                self.current_para = self.doc.add_paragraph()
                self.current_para.style = self.current_style
            elif tag == "p":
                self.current_style = "Normal"
                self.current_para = self.doc.add_paragraph()
            elif tag in ("ul", "ol"):
                self.list_type = "List Bullet"
                self.list_level += 1
            elif tag == "li":
                self.current_para = self.doc.add_paragraph(style="List Bullet")
            elif tag == "strong" or tag == "b":
                self.bold = True
            elif tag == "em" or tag == "i":
                self.italic = True

        def handle_endtag(self, tag):
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6", "p", "li"):
                self.current_para = None
            elif tag in ("ul", "ol"):
                self.list_type = None
                self.list_level = max(0, self.list_level - 1)
            elif tag == "strong" or tag == "b":
                self.bold = False
            elif tag == "em" or tag == "i":
                self.italic = False

        def handle_data(self, data):
            if not data.strip():
                return
            if self.current_para is None:
                self.current_para = self.doc.add_paragraph()
            run = self.current_para.add_run(data)
            run.bold = self.bold
            run.italic = self.italic

    # Method to process a document set and take out the text from the document then pass it to the LLM client to make the final summary
    async def process_document_set(
        self, doc_set: Dict[str, Any], input_dir: Path
    ) -> Dict[str, Any]:
        """Process a document set and generate a summary."""
        try:
            set_name = doc_set["name"]
            documents = doc_set["documents"]
            document_analyses = []

            for doc_name in documents:
                doc_path = input_dir / doc_name
                try:
                    doc = Document(doc_path)
                    text_content = "\n".join(
                        [para.text for para in doc.paragraphs if para.text.strip()]
                    )
                    doc_analysis = await self.analyze_document(text_content, doc_name)
                    document_analyses.append(doc_analysis)
                except Exception as e:
                    document_analyses.append(
                        self._handle_error("processing document", doc_name, e)
                    )

            set_summary = await self.generate_comprehensive_summary(document_analyses)
            return {
                "set_name": set_name,
                "summary": set_summary,
                "documents": documents,
            }
        except Exception as e:
            logger.error(f"Error processing document set: {str(e)}")
            return {
                "set_name": doc_set.get("name", "Unknown"),
                "summary": f"Error processing document set: {str(e)}",
                "documents": doc_set.get("documents", []),
            }

    # Method to split document into overlapping chunks
    def chunk_document(self, text: str) -> List[str]:
        """Split document into overlapping chunks."""
        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            # Calculate end position for this chunk
            end = start + self.chunk_size

            # If this is not the first chunk, include some overlap from previous chunk
            if start > 0:
                start = start - self.chunk_overlap

            # If this is the last chunk, take all remaining text
            if end >= text_length:
                chunks.append(text[start:])
                break

            # Find the last period or newline in the chunk to break at a natural point
            last_period = text.rfind(".", start, end)
            last_newline = text.rfind("\n", start, end)
            break_point = max(last_period, last_newline)

            if break_point > start:
                end = break_point + 1

            chunks.append(text[start:end])
            start = end

        return chunks

    # Method to analyze a document
    async def analyze_document(self, text: str, doc_name: str) -> Dict[str, Any]:
        """Analyze a document and extract key information."""
        try:
            # Split document into chunks
            chunks = self.chunk_document(text)
            chunk_analyses = []

            # Process each chunk
            for i, chunk in enumerate(chunks, 1):
                prompt = f"""
                Analyze this part of the document (Part {i} of {len(chunks)}) and provide:
                1. A concise summary of the main content
                2. Any critical or important information that needs special attention

                Document: {doc_name}
                Content: {chunk}
                """

                chunk_analysis = await self.llm_client.generate_content(prompt)
                if chunk_analysis:
                    chunk_analyses.append(chunk_analysis)

            # Combine chunk analyses
            if not chunk_analyses:
                return {
                    "document_name": doc_name,
                    "analysis": "No analysis available - document may be empty or unreadable",
                }

            # Create final summary of all chunks
            summary_prompt = f"""
            Create a comprehensive analysis of this document by combining the analyses of its parts.
            Focus on two main aspects:
            1. Executive Summary: Provide a clear, concise summary of the entire document
            2. Important Information: List any critical points, key findings, or information that requires special attention

            Document: {doc_name}
            Part Analyses:
            {chr(10).join([f"Part {i+1}: {analysis}" for i, analysis in enumerate(chunk_analyses)])}
            """

            final_analysis = await self.llm_client.generate_content(summary_prompt)

            return {
                "document_name": doc_name,
                "analysis": final_analysis or "No analysis available",
            }

        except Exception as e:
            logger.error(f"Error analyzing document {doc_name}: {str(e)}")
            return {
                "document_name": doc_name,
                "analysis": f"Error analyzing document: {str(e)}",
            }

    # Generate a focused summary
    async def generate_comprehensive_summary(
        self, document_analyses: List[Dict[str, Any]]
    ) -> str:
        """Generate a focused summary for a pair of documents."""
        try:
            if not document_analyses:
                return "No documents available for analysis"

            # Create a focused prompt for the LLM
            prompt = f"""
            Create a concise summary of these related documents, focusing on:
            1. Main Points
            2. Document Summaries
            3. Key Findings
            4. Important Information

            Documents:
            {chr(10).join([f"- {analysis['document_name']}: {analysis.get('analysis', 'No analysis available')}" for analysis in document_analyses])}
            """

            summary = await self.llm_client.generate_content(prompt)

            if not summary:
                return "Unable to generate summary - please check document contents"

            return summary
        except Exception as e:
            logger.error(f"Error generating summary: {str(e)}")
            return f"Error generating summary: {str(e)}"

    # Method to format text by converting markdown to Word formatting using markdown2 and HTMLParser
    def _format_text(self, text: str, paragraph) -> None:
        """Format text by converting markdown to Word formatting using markdown2 and HTMLParser."""
        try:
            doc = paragraph._parent
            html = markdown2.markdown(text)
            style_map = {
                "h1": "SectionHeader",
                "h2": "Heading 3",
                "h3": "Heading 4",
                "h4": "Heading 5",
                "h5": "Heading 6",
                "h6": "Heading 6",
            }
            parser = self.DocxHTMLParser(doc, style_map)
            parser.feed(html)
        except Exception as e:
            logger.error(f"Error formatting text: {str(e)}")
            paragraph.add_run(text)

    # Create a context document from processed document sets
    def create_context_document(
        self, document_sets: List[Dict[str, Any]], input_dir: Path, output_file: str
    ) -> Document:
        """Create a context document from processed document sets."""
        try:
            doc = Document()
            self.style_manager.init_document_styles(doc)
            title = doc.add_paragraph("Merged Document")
            title.style = self.style_manager.get_safe_style(doc, "CustomTitle")
            for set_info in document_sets:
                try:
                    header = doc.add_paragraph(set_info["set_name"])
                    header.style = self.style_manager.get_safe_style(doc, "SectionHeader")
                    if set_info.get("summary"):
                        summary_para = doc.add_paragraph()
                        self._format_text(set_info["summary"], summary_para)
                        summary_para.style = self.style_manager.get_safe_style(doc, "Summary")
                    docs_header = doc.add_paragraph("Documents Analyzed")
                    docs_header.style = self.style_manager.get_safe_style(doc, "SectionHeader")
                    for doc_path in set_info.get("documents", []):
                        doc_name = Path(doc_path).name
                        bullet_para = doc.add_paragraph(style="List Bullet")
                        bullet_para.add_run(doc_name)
                    doc.add_paragraph()
                except Exception as e:
                    logger.error(f"Error processing set {set_info.get('set_name', 'Unknown')}: {str(e)}")
                    continue
            return doc
        except Exception as e:
            logger.error(f"Error creating context document: {str(e)}")
            raise
